#!/usr/bin/env python3
"""
LLM-Based Co-Document Analyzer
Compares two related documents to detect fraud, inconsistencies, and discrepancies

Use Cases:
- Bill vs Supporting Documents (medical bills vs prescriptions, lab reports)
- Salary Slip vs Employment Documents (offer letter, appointment letter)
- Invoice vs Purchase Order
- Agreement vs Supporting Evidence
- Identity Documents Cross-Verification

Author: AI Fraud Detection System
Date: November 6, 2025
"""

import os
import sys
import json
import argparse
import datetime
import base64
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
import re

import fitz  # PyMuPDF
from PIL import Image

# Add parent directory to path for config import
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from config import AzureOpenAIConfig, get_openai_client
from prompt_manager import get_co_document_comparison_prompt, get_co_document_system_prompt

# Initialize Azure OpenAI client from environment
client = get_openai_client()


class CoDocumentAnalyzer:
    """LLM-based analyzer for comparing two related documents"""
    
    def __init__(self):
        self.client = client
        self.report_dir = "co_document_reports"
        os.makedirs(self.report_dir, exist_ok=True)
    
    def encode_image(self, image_path: str) -> str:
        """Encode image to base64"""
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    
    def extract_text_from_pdf(self, pdf_path: str) -> Tuple[str, List[str]]:
        """Extract text and images from PDF - converts ALL pages to images"""
        doc = fitz.open(pdf_path)
        all_text = ""
        image_paths = []
        
        temp_dir = "temp_extracted_images"
        os.makedirs(temp_dir, exist_ok=True)
        
        pdf_name = Path(pdf_path).stem
        
        # Extract text and convert each page to image
        for page_num, page in enumerate(doc):
            all_text += page.get_text()
            
            # Convert each page to high-quality image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
            temp_img = os.path.join(temp_dir, f"{pdf_name}_page_{page_num}.png")
            pix.save(temp_img)
            image_paths.append(temp_img)
        
        doc.close()
        return all_text, image_paths
    
    def convert_docx_to_images(self, docx_path: str) -> List[str]:
        """Convert DOCX to images using LLM vision approach"""
        try:
            # Try using python-docx and Pillow to render
            from docx import Document
            from docx.shared import Inches
            import io
            
            doc = Document(docx_path)
            temp_dir = "temp_extracted_images"
            os.makedirs(temp_dir, exist_ok=True)
            
            # For now, we'll use a simpler approach: convert via LibreOffice or similar
            # If not available, create a placeholder
            print("   Converting DOCX to image for vision analysis...")
            
            # Try using LibreOffice if available
            docx_name = Path(docx_path).stem
            output_pdf = os.path.join(temp_dir, f"{docx_name}_converted.pdf")
            
            # Attempt conversion using soffice (LibreOffice)
            import subprocess
            try:
                subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', temp_dir,
                    docx_path
                ], check=True, timeout=30, capture_output=True)
                
                # Now convert the PDF to images
                if os.path.exists(output_pdf):
                    _, images = self.extract_text_from_pdf(output_pdf)
                    return images
            except:
                pass
            
            # Fallback: Just use vision on a screenshot/placeholder
            # Create a simple text extraction for context
            doc = Document(docx_path)
            text_content = '\n'.join([para.text for para in doc.paragraphs])
            
            # Create a simple image with text (fallback)
            from PIL import Image as PILImage, ImageDraw, ImageFont
            
            # Create a white background image
            img_width, img_height = 1200, 1600
            img = PILImage.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use a default font
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
            
            # Draw text on image
            y_position = 50
            for line in text_content.split('\n')[:50]:  # First 50 lines
                if line.strip():
                    # Wrap long lines
                    words = line.split()
                    current_line = ""
                    for word in words:
                        test_line = current_line + word + " "
                        if len(test_line) * 10 < img_width - 100:
                            current_line = test_line
                        else:
                            if current_line:
                                draw.text((50, y_position), current_line, fill='black', font=font)
                                y_position += 30
                            current_line = word + " "
                    if current_line:
                        draw.text((50, y_position), current_line, fill='black', font=font)
                        y_position += 30
                
                if y_position > img_height - 50:
                    break
            
            # Save the image
            output_image = os.path.join(temp_dir, f"{docx_name}_page_0.png")
            img.save(output_image)
            
            return [output_image]
            
        except Exception as e:
            print(f"   âš  Could not convert DOCX: {e}")
            print(f"   Will use vision-only analysis")
            # Return empty list to trigger vision-only mode
            return []

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and extract relevant information"""
        
        print(f"\nðŸ“„ Processing: {file_path}")
        file_ext = Path(file_path).suffix.lower()
        
        doc_info = {
            'file_path': file_path,
            'file_name': Path(file_path).name,
            'file_type': file_ext,
            'text': '',
            'image_paths': [],
            'metadata': {}
        }
        
        if file_ext == '.pdf':
            print("   Extracting from PDF...")
            text, images = self.extract_text_from_pdf(file_path)
            doc_info['text'] = text
            doc_info['image_paths'] = images
            print(f"   âœ“ Extracted {len(text)} characters, {len(images)} images")
        
        elif file_ext in ['.docx', '.doc']:
            print("   Processing Word document...")
            # Extract text for context
            try:
                from docx import Document
                doc = Document(file_path)
                doc_info['text'] = '\n'.join([para.text for para in doc.paragraphs])
                print(f"   âœ“ Extracted {len(doc_info['text'])} characters")
            except Exception as e:
                print(f"   âš  Could not extract text: {e}")
            
            # Convert to images for vision analysis
            images = self.convert_docx_to_images(file_path)
            if images:
                doc_info['image_paths'] = images
                print(f"   âœ“ Converted to {len(images)} images for vision analysis")
            else:
                # Vision-only mode - create a note
                print(f"   â„¹ Using vision-only mode (provide image manually if needed)")
        
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
            print("   Processing image...")
            doc_info['image_paths'] = [file_path]
            
            # Try OCR if available
            try:
                import pytesseract
                from PIL import Image as PILImage
                img = PILImage.open(file_path)
                doc_info['text'] = pytesseract.image_to_string(img)
                print(f"   âœ“ Extracted {len(doc_info['text'])} characters via OCR")
            except:
                print("   â„¹ Using vision-only mode")
        
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: PDF, DOCX, JPG, PNG, BMP")
        
        return doc_info
    
    def create_comparison_prompt(self, doc1_type: str = "Document 1", doc2_type: str = "Document 2") -> str:
        """Create comprehensive prompt for comparing two documents"""
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")
        return get_co_document_comparison_prompt(doc1_type, doc2_type, current_date)
    
    def analyze_co_documents(
        self, 
        doc1_info: Dict[str, Any], 
        doc2_info: Dict[str, Any],
        doc1_type: str = "Document 1",
        doc2_type: str = "Document 2"
    ) -> Dict[str, Any]:
        """Analyze two documents together using LLM reasoning"""
        
        print(f"\nðŸ¤– Analyzing documents with AI...")
        print(f"   Comparing: {doc1_type} vs {doc2_type}")
        
        # Check if we have images
        if not doc1_info['image_paths'] and not doc2_info['image_paths']:
            return self.analyze_text_only(doc1_info, doc2_info, doc1_type, doc2_type)
        
        if not doc1_info['image_paths'] or not doc2_info['image_paths']:
            if not doc1_info['image_paths'] and not doc1_info['text']:
                return {
                    "verdict": "ERROR",
                    "error": f"Could not extract content from Document 1 ({doc1_info['file_name']})",
                    "confidence_score": 0
                }
            if not doc2_info['image_paths'] and not doc2_info['text']:
                return {
                    "verdict": "ERROR",
                    "error": f"Could not extract content from Document 2 ({doc2_info['file_name']})",
                    "confidence_score": 0
                }
        
        # Build content array with all pages
        content = []
        
        # Create prompt
        prompt = self.create_comparison_prompt(doc1_type, doc2_type)
        
        # Add extracted text context
        context = f"\n\n**EXTRACTED TEXT - {doc1_type.upper()}:**\n{doc1_info['text'][:5000]}\n\n"
        context += f"**EXTRACTED TEXT - {doc2_type.upper()}:**\n{doc2_info['text'][:5000]}"
        
        content.append({
            "type": "text",
            "text": prompt + context
        })
        
        # Add all pages from document 1
        for i, img_path in enumerate(doc1_info['image_paths']):
            base64_image = self.encode_image(img_path)
            content.append({
                "type": "text",
                "text": f"\n**{doc1_type.upper()} - PAGE {i+1}:**"
            })
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{base64_image}"
                }
            })
        
        # Add all pages from document 2
        for i, img_path in enumerate(doc2_info['image_paths']):
            base64_image = self.encode_image(img_path)
            content.append({
                "type": "text",
                "text": f"\n**{doc2_type.upper()} - PAGE {i+1}:**"
            })
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{base64_image}"
                }
            })
        
        messages = [
            {
                "role": "system",
                "content": "You are an expert fraud detection AI specializing in cross-document analysis and fraud pattern recognition."
            },
            {
                "role": "user",
                "content": content
            }
        ]
        
        try:
            print("   Sending to Azure OpenAI for analysis...")
            response = self.client.chat.completions.create(
                model=AzureOpenAIConfig.DEPLOYMENT,
                messages=messages,
                max_completion_tokens=6000,
                temperature=0.1  # Low temperature for consistent analysis
            )
            
            response_text = response.choices[0].message.content
            print("   âœ“ Analysis complete")
            
            # Try to extract JSON from response
            json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                try:
                    analysis = json.loads(json_match.group(1))
                except json.JSONDecodeError as e:
                    print(f"   âš  JSON parsing error: {e}")
                    print(f"   Attempting to fix common JSON issues...")
                    # Try to fix common JSON issues
                    json_text = json_match.group(1)
                    # Remove trailing commas
                    json_text = re.sub(r',\s*}', '}', json_text)
                    json_text = re.sub(r',\s*]', ']', json_text)
                    try:
                        analysis = json.loads(json_text)
                    except:
                        analysis = {
                            "verdict": "ERROR",
                            "confidence_score": 0,
                            "raw_response": response_text[:2000],
                            "error": f"JSON parsing failed: {str(e)}"
                        }
            else:
                # Try to parse entire response as JSON
                try:
                    analysis = json.loads(response_text)
                except:
                    analysis = {
                        "verdict": "ERROR",
                        "confidence_score": 0,
                        "raw_response": response_text[:2000],
                        "error": "Could not parse JSON from response"
                    }
            
            # Add metadata
            analysis['metadata'] = {
                'document_1': {
                    'file': doc1_info['file_name'],
                    'type': doc1_type,
                    'path': doc1_info['file_path']
                },
                'document_2': {
                    'file': doc2_info['file_name'],
                    'type': doc2_type,
                    'path': doc2_info['file_path']
                },
                'analysis_timestamp': datetime.datetime.now().isoformat(),
                'model': AzureOpenAIConfig.DEPLOYMENT
            }
            
            return analysis
            
        except Exception as e:
            return {
                "verdict": "ERROR",
                "error": str(e),
                "confidence_score": 0,
                "metadata": {
                    "error_details": str(e)
                }
            }
    
    def analyze_text_only(
        self,
        doc1_info: Dict[str, Any],
        doc2_info: Dict[str, Any],
        doc1_type: str = "Document 1",
        doc2_type: str = "Document 2"
    ) -> Dict[str, Any]:
        """Analyze two documents using text-only (no images available)"""
        
        print(f"   Using text-only analysis mode...")
        
        # Create prompt for text-only analysis
        prompt = self.create_comparison_prompt(doc1_type, doc2_type)
        
        # Add text content
        context = f"\n\n**FULL TEXT - {doc1_type.upper()}:**\n{doc1_info['text'][:5000]}\n\n"
        context += f"**FULL TEXT - {doc2_type.upper()}:**\n{doc2_info['text'][:5000]}"
        context += "\n\n**NOTE: No images available - analyzing text content only**"
        
        messages = [
            {
                "role": "system",
                "content": "You are an expert fraud detection AI specializing in cross-document analysis and fraud pattern recognition."
            },
            {
                "role": "user",
                "content": prompt + context
            }
        ]
        
        try:
            print("   Sending to Azure OpenAI for analysis...")
            response = self.client.chat.completions.create(
                model=AzureOpenAIConfig.DEPLOYMENT,
                messages=messages,
                max_completion_tokens=6000,
                temperature=0.1
            )
            
            response_text = response.choices[0].message.content
            print("   âœ“ Analysis complete")
            
            # Parse JSON response
            json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                try:
                    analysis = json.loads(json_match.group(1))
                except json.JSONDecodeError as e:
                    print(f"   âš  JSON parsing error: {e}")
                    # Try to fix common JSON issues
                    json_text = json_match.group(1)
                    json_text = re.sub(r',\s*}', '}', json_text)
                    json_text = re.sub(r',\s*]', ']', json_text)
                    try:
                        analysis = json.loads(json_text)
                    except:
                        analysis = {
                            "verdict": "ERROR",
                            "confidence_score": 0,
                            "raw_response": response_text[:2000],
                            "error": f"JSON parsing failed: {str(e)}"
                        }
            else:
                try:
                    analysis = json.loads(response_text)
                except:
                    analysis = {
                        "verdict": "ERROR",
                        "confidence_score": 0,
                        "raw_response": response_text[:2000],
                        "error": "Could not parse JSON from response"
                    }
            
            # Add metadata
            analysis['metadata'] = {
                'document_1': {
                    'file': doc1_info['file_name'],
                    'type': doc1_type,
                    'path': doc1_info['file_path']
                },
                'document_2': {
                    'file': doc2_info['file_name'],
                    'type': doc2_type,
                    'path': doc2_info['file_path']
                },
                'analysis_timestamp': datetime.datetime.now().isoformat(),
                'model': AzureOpenAIConfig.DEPLOYMENT,
                'mode': 'text_only'
            }
            
            return analysis
            
        except Exception as e:
            return {
                "verdict": "ERROR",
                "error": str(e),
                "confidence_score": 0,
                "metadata": {
                    "error_details": str(e)
                }
            }
    
    def display_results(self, analysis: Dict[str, Any]):
        """Display analysis results in a clear, professional format"""
        
        print(f"\n{'='*80}")
        print("ANALYSIS REPORT")
        print(f"{'='*80}\n")
        
        if analysis.get('verdict') == 'ERROR':
            print(f"âŒ ERROR: {analysis.get('error', 'Unknown error')}")
            if 'raw_response' in analysis:
                print(f"\nRaw Response:\n{analysis['raw_response'][:1000]}...")
            return
        
        # Verdict with emoji
        verdict = analysis.get('verdict', 'UNKNOWN')
        verdict_emoji = {
            'CONSISTENT': 'âœ…',
            'INCONSISTENT': 'âš ï¸',
            'SUSPICIOUS': 'ðŸŸ ',
            'FRAUDULENT': 'ðŸš¨',
            'UNKNOWN': 'â“'
        }
        
        print(f"VERDICT: {verdict} {verdict_emoji.get(verdict, 'â“')}")
        print(f"CONFIDENCE: {analysis.get('confidence_score', 0)}%")
        print(f"RISK LEVEL: {analysis.get('risk_level', 'UNKNOWN')}")
        
        if 'summary' in analysis:
            print(f"\nEXECUTIVE SUMMARY:")
            print(f"{analysis['summary']}")
        
        # Document Analysis - Simplified
        if 'document_analysis' in analysis:
            print(f"\n{'='*80}")
            print("DOCUMENT DETAILS:")
            print(f"{'='*80}")
            
            for doc_key in ['document_1', 'document_2']:
                if doc_key in analysis['document_analysis']:
                    doc = analysis['document_analysis'][doc_key]
                    doc_num = doc_key.split('_')[1]
                    print(f"\n[{doc_num.upper()}] {doc.get('type', 'Unknown')}")
                    
                    if 'key_details' in doc:
                        details = doc['key_details']
                        if details.get('name'):
                            print(f"  Name: {details['name']}")
                        if details.get('date'):
                            print(f"  Date: {details['date']}")
                        if details.get('amount'):
                            print(f"  Amount: {details['amount']}")
                    
                    print(f"  Quality: {doc.get('quality_assessment', 'N/A')}")
        
        # Comparison Results - Concise
        if 'comparison_results' in analysis:
            print(f"\n{'='*80}")
            print("KEY FINDINGS:")
            print(f"{'='*80}")
            
            comp = analysis['comparison_results']
            checks = []
            
            # Identity Matching
            if 'identity_matching' in comp:
                identity = comp['identity_matching']
                status = 'âœ“' if identity.get('verdict') == 'PASS' else 'âœ—'
                checks.append(f"{status} Identity Matching: {identity.get('verdict', 'N/A')}")
            
            # Date Consistency
            if 'date_consistency' in comp:
                dates = comp['date_consistency']
                status = 'âœ“' if dates.get('verdict') == 'PASS' else 'âœ—'
                checks.append(f"{status} Date Consistency: {dates.get('verdict', 'N/A')}")
            
            # Financial Reconciliation
            if 'financial_reconciliation' in comp:
                finance = comp['financial_reconciliation']
                if finance.get('status') != 'NOT_APPLICABLE':
                    status = 'âœ“' if finance.get('verdict') == 'PASS' else 'âœ—'
                    checks.append(f"{status} Financial Match: {finance.get('verdict', 'N/A')}")
            
            # Content Consistency
            if 'content_consistency' in comp:
                content = comp['content_consistency']
                status = 'âœ“' if content.get('verdict') == 'PASS' else 'âœ—'
                checks.append(f"{status} Content Match: {content.get('verdict', 'N/A')}")
            
            for check in checks:
                print(f"  {check}")
        
        # Red Flags - Top issues only
        if 'red_flags' in analysis and analysis['red_flags']:
            print(f"\n{'='*80}")
            print("CRITICAL ISSUES:")
            print(f"{'='*80}")
            
            # Show only HIGH and CRITICAL flags
            critical_flags = [f for f in analysis['red_flags'] if f.get('severity') in ['HIGH', 'CRITICAL']]
            
            if critical_flags:
                for i, flag in enumerate(critical_flags[:5], 1):  # Top 5 critical issues
                    severity = flag.get('severity', 'UNKNOWN')
                    emoji = 'ðŸ”´' if severity in ['HIGH', 'CRITICAL'] else 'ðŸŸ '
                    
                    print(f"\n{emoji} [{severity}] {flag.get('description', 'No description')}")
                    
                    if 'evidence' in flag and isinstance(flag['evidence'], dict):
                        if flag['evidence'].get('discrepancy'):
                            print(f"   Issue: {flag['evidence']['discrepancy']}")
            else:
                # Show top 3 flags if no critical ones
                for i, flag in enumerate(analysis['red_flags'][:3], 1):
                    print(f"\n{i}. [{flag.get('severity', 'N/A')}] {flag.get('description', 'No description')}")
        
        # Fraud Indicators - Simplified
        if 'fraud_indicators' in analysis and analysis['fraud_indicators']:
            print(f"\n{'='*80}")
            print("FRAUD RISK ASSESSMENT:")
            print(f"{'='*80}")
            
            for indicator in analysis['fraud_indicators'][:3]:  # Top 3 only
                conf = indicator.get('confidence', 0)
                if conf > 40:  # Only show significant risks
                    print(f"\nâ€¢ {indicator.get('type', 'Unknown').replace('_', ' ').title()} ({conf}% confidence)")
                    print(f"  {indicator.get('description', 'No description')}")
        
        # Recommendations - Concise
        if 'recommendations' in analysis and analysis['recommendations']:
            print(f"\n{'='*80}")
            print("RECOMMENDED ACTIONS:")
            print(f"{'='*80}")
            for i, rec in enumerate(analysis['recommendations'][:5], 1):  # Top 5
                print(f"{i}. {rec}")
    
    def save_report(self, analysis: Dict[str, Any]) -> str:
        """Save analysis report to file"""
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create filename from document names
        base_name = "co_document_analysis"
        if 'metadata' in analysis and 'document_1' in analysis.get('metadata', {}):
            try:
                doc1_name = Path(analysis['metadata']['document_1']['file']).stem
                doc2_name = Path(analysis['metadata']['document_2']['file']).stem
                base_name = f"{doc1_name}_vs_{doc2_name}"
            except (KeyError, TypeError):
                pass
        
        # Save JSON report
        json_report = os.path.join(self.report_dir, f"{base_name}_{timestamp}.json")
        with open(json_report, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, indent=2, ensure_ascii=False)
        
        # Save text report
        txt_report = os.path.join(self.report_dir, f"{base_name}_{timestamp}.txt")
        with open(txt_report, 'w', encoding='utf-8') as f:
            f.write(f"CO-DOCUMENT ANALYSIS REPORT\n")
            f.write(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"{'='*80}\n\n")
            
            if 'metadata' in analysis:
                f.write(f"Document 1: {analysis['metadata']['document_1']['file']}\n")
                f.write(f"Document 2: {analysis['metadata']['document_2']['file']}\n\n")
            
            f.write(f"VERDICT: {analysis.get('verdict', 'UNKNOWN')}\n")
            f.write(f"CONFIDENCE: {analysis.get('confidence_score', 0)}%\n")
            f.write(f"RISK LEVEL: {analysis.get('risk_level', 'UNKNOWN')}\n\n")
            
            if 'summary' in analysis:
                f.write(f"SUMMARY:\n{analysis['summary']}\n\n")
            
            f.write(f"{'='*80}\n")
            f.write("FULL ANALYSIS (JSON):\n")
            f.write(f"{'='*80}\n")
            f.write(json.dumps(analysis, indent=2, ensure_ascii=False))
        
        return json_report
    
    def compare_documents(
        self,
        file1: str,
        file2: str,
        doc1_type: str = "Document 1",
        doc2_type: str = "Document 2"
    ) -> Dict[str, Any]:
        """Main function to compare two documents"""
        
        print(f"\n{'='*80}")
        print("CO-DOCUMENT FRAUD ANALYSIS")
        print(f"{'='*80}")
        print(f"\nðŸ” Analyzing relationship between:")
        print(f"   1. {doc1_type}: {file1}")
        print(f"   2. {doc2_type}: {file2}")
        print(f"{'='*80}")
        
        # Process both documents
        try:
            doc1_info = self.process_document(file1)
            doc2_info = self.process_document(file2)
        except Exception as e:
            print(f"\nâŒ Error processing documents: {e}")
            return {"verdict": "ERROR", "error": str(e)}
        
        # Analyze together
        analysis = self.analyze_co_documents(doc1_info, doc2_info, doc1_type, doc2_type)
        
        # Display results
        self.display_results(analysis)
        
        # Save report (skip if error)
        if analysis.get('verdict') != 'ERROR':
            report_path = self.save_report(analysis)
            print(f"\nðŸ“Š Report saved: {report_path}")
        else:
            print(f"\nâš  Report not saved due to analysis error")
        
        return analysis


def main():
    """Main entry point"""
    
    parser = argparse.ArgumentParser(
        description="Compare two related documents for fraud detection and inconsistencies",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
EXAMPLES:

1. Compare medical bill with prescription:
   python codocanalyser.py bill.pdf prescription.pdf --type1 "Medical Bill" --type2 "Prescription"

2. Compare salary slip with offer letter:
   python codocanalyser.py salary_slip.pdf offer_letter.pdf --type1 "Salary Slip" --type2 "Offer Letter"

3. Compare invoice with purchase order:
   python codocanalyser.py invoice.pdf po.pdf --type1 "Invoice" --type2 "Purchase Order"

4. Simple comparison (auto-detect types):
   python codocanalyser.py doc1.pdf doc2.pdf

DOCUMENT TYPE EXAMPLES:
- Medical: "Medical Bill", "Prescription", "Lab Report", "Discharge Summary"
- Employment: "Salary Slip", "Offer Letter", "Appointment Letter", "Experience Certificate"
- Business: "Invoice", "Purchase Order", "Delivery Note", "Agreement"
- Identity: "PAN Card", "Aadhaar Card", "Passport", "Driver's License"
        """
    )
    
    parser.add_argument('file1', help='Path to first document (PDF, JPG, PNG)')
    parser.add_argument('file2', help='Path to second document (PDF, JPG, PNG)')
    parser.add_argument('--type1', default='Document 1', help='Type/description of first document')
    parser.add_argument('--type2', default='Document 2', help='Type/description of second document')
    
    args = parser.parse_args()
    
    # Validate files exist
    if not os.path.exists(args.file1):
        print(f"âŒ Error: File not found: {args.file1}")
        sys.exit(1)
    
    if not os.path.exists(args.file2):
        print(f"âŒ Error: File not found: {args.file2}")
        sys.exit(1)
    
    # Run comparison
    analyzer = CoDocumentAnalyzer()
    analyzer.compare_documents(args.file1, args.file2, args.type1, args.type2)


if __name__ == "__main__":
    main()
